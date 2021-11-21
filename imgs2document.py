#The application is used to create a pdf or docx file from a folder of images, 
#images are appended to the pdf or word file in the order of the filenames

import easygui
import os
import PIL
import fstring
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Inches,Mm,Cm
from fpdf import FPDF
from tqdm import tqdm
import PySimpleGUI as sg
import datetime

#import a string similarity function to compare the output file name extension aginst the output format to see which is the most similar            
from difflib import SequenceMatcher
def similarity(a, b):
    return SequenceMatcher(None, a, b).ratio()

def mainGUI():
    #create the GUI, including file path and output file name, and the orentation of the page layout in word and pdf
    #file path has a button to open directory dialog
    layout = [[sg.Text('Please select the directory of images')],
              [sg.Input(key='-INPUT-', enable_events=True, size=(50, 1)), sg.FolderBrowse(target='-INPUT-')],
              [sg.Text('Please input the output file name and format')],
              #add a radio group to select the output format
              [sg.Input(key='-OUTPUT-', size=(50, 1))],
              [sg.Radio('PDF', "FORMAT", default=True, enable_events=True, key='PDF'), sg.Radio('Word', "FORMAT", enable_events=True,key='Word')],              
              [sg.Text('Please select the orientation of the page layout in word and pdf')],
              [sg.Radio('Portrait', "ORIENTATION", default=True, key='-PORTRAIT-'), sg.Radio('Landscape', "ORIENTATION", key='-LANDSCAPE-')],
              [sg.Submit(), sg.Cancel()]]
    window = sg.Window('Images to document', layout)    

    while True:        
        event, values = window.read()
        #debug print current time, the event and values
        #print(f"{datetime.datetime.now()} {event} {values}")
        print(f"[ {datetime.datetime.now()} ] :  event={event}, values={values}")        
        #when the user select a folder, the input field will be updated with the folder path, in the meantime, 
        # the output file name will be updated with the folder name plus .pdf or .docx extension, but to use which one, is decided by the output format
        if event in ['-INPUT-', 'PDF','Word']:                       
            #if the -INPUT- is not empty and if -INPUT- is a valid folder, update the -OUTPUT- 
            # and the output file name will be updated with the folder name plus .pdf or .docx extension, but to use which one, is decided by the output format

            if values['-INPUT-'] != '' and os.path.isdir(values['-INPUT-']):
                #get the folder name
                folder_name = values['-INPUT-']
                #get the output file name
                if values['PDF']:
                    file_name = folder_name + '.pdf'
                else:
                    file_name = folder_name + '.docx'
                #update the -OUTPUT-
                window['-OUTPUT-'].update(file_name)
            
        #if user select cancel, return None
        if event in (None, 'Cancel'):
            window.close()
            return None

        #if user select submit, return the input values  
        if event == 'Submit':
            window.close()
            return values

def main():
    #get the input from the GUI
    values = mainGUI()
    if values is None:
        return
    #get the path of the folder
    image_folder = values['-INPUT-']
    #get the name of the output file
    file_name = values['-OUTPUT-']    

    #get the page layout
    if values['-PORTRAIT-']:
        page_layout = 'P'
    else:
        page_layout = 'L'
    
    #check if the folder is valid
    if not os.path.isdir(image_folder):
        print(f"{image_folder} is not a valid folder")
        return
    
    #check if the output file name is valid
    if not file_name.endswith(('.pdf', '.docx')):
        print(f"{file_name} is not a valid file name")
        return
    else:
        #check if the file exists, delete it if it does
        if os.path.exists(file_name):
            os.remove(file_name)
    
    #get the output format
    if values['PDF']:        
        count = insertImages2PDF(img_path=image_folder, outputfile=file_name, page_layout=page_layout,windowGUI=True)
    else:
        count = insertImages2WordDoc(img_path=image_folder, outputfile=file_name, page_layout=page_layout,windowGUI=True)  

def insertImages2WordDoc(img_path=None, outputfile=None, page_layout='P',windowGUI=False):
    # Get the path to the folder
    if img_path is None:
        image_folder = easygui.diropenbox()
    else:
        image_folder = img_path
    
    # generate the name of the pdf file based on path
    if outputfile is None:
        doc_name = os.path.basename(image_folder) + ".docx"
    else:
        doc_name = outputfile
    
    #check if there is a pdf file with the same name in the folder, delete it if there is
    if os.path.isfile(doc_name):
        os.remove(doc_name)
        
    image_list=[]
    
    #walk through the folder and get all the images    
    for root, dirs, files in os.walk(image_folder):
        for file in files:
            #first conver filaname to lowercase
            file_name = file.lower()
            if file_name.endswith((".jpg", ".png", ".jpeg")):
                #test if the file is really a picture by reading the first few bytes, if this is not a picture, skip it
                try:
                    PIL.Image.open(os.path.join(root, file))
                    image_list.append(os.path.join(root, file))
                except Exception as e:
                    if windowGUI:
                        sg.popup_error(f"[Error] : inserting {file} failed, skip! Failure reason: \n Exception : {e}\n") 
                    else:                        
                        print(f"[Error] : inserting {file} failed, skip! Failure reason: \n Exception : {e}\n")
            
    #sort the list of images
    image_list.sort()
    print(f"{len(image_list)} images found in the folder {image_folder}:")        
    
    #now we have a list of images, we can create the pdf file
    doc = docx.Document()      
    count=0        
    #setup the docx page orientation based on the input page_layout      
    if page_layout == 'P':
        doc.sections[0].orientation = WD_ORIENT.PORTRAIT
        #setup page size to A4 Portrait
        doc.sections[0].page_width = Cm(21.0)
        doc.sections[0].page_height = Cm(29.7)
    else:
        doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
        #setup page size to A4 Landscape
        doc.sections[0].page_width = Cm(29.7)
        doc.sections[0].page_height = Cm(21.0)           
    for image in tqdm(image_list, desc="Inserting images to document",total=len(image_list)):
       #insert the image to doc, make sure the size of the image is not bigger than the page and keep ratio           
      #in case of any exception, skip the image
        try:                
            p=doc.add_paragraph()
            p.alignment=docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run()
            r.add_picture(image, width=docx.shared.Inches(6))
            doc.add_page_break() 
            count+=1                                           
        except Exception as e:
            if windowGUI:
                sg.popup_error(f"[Exception]: when handling file {image}, Exception happened:\n{e}\n")                        
            else:
                print(f"[Exception]: when handling file {image}, Exception happened:\n{e}\n")                                   

    doc.save(doc_name)        
    if windowGUI:        
        sg.popup(f"Word file {doc_name} created, {count} images added!")        
    else:
        print(f"Word file {doc_name} created, {count} images added!")    
    return count

def insertImages2PDF(img_path=None, outputfile=None,page_layout='P',windowGUI=False):
    
    # Get the path to the folder
    if img_path is None:
        image_folder = easygui.diropenbox()
    else:
        image_folder = img_path
    
    # generate the name of the pdf file based on path
    if outputfile is None:
        pdf_name = os.path.basename(image_folder) + ".pdf"
    else:
        pdf_name = outputfile
    
    #check if there is a pdf file with the same name in the folder, delete it if there is
    if os.path.isfile(pdf_name):
        os.remove(pdf_name)
        
    image_list=[]    
    
    #walk through the folder and get all the images    
    for root, dirs, files in os.walk(image_folder):
        for file in files:
            #first conver filaname to lowercase
            file_name = file.lower()            
            if file_name.endswith((".jpg", ".png", ".jpeg")):
                image_list.append(os.path.join(root, file))
            
    #sort the list of images
    image_list.sort()
    if not windowGUI:
        print(f"{len(image_list)} images found in the folder {image_folder}:")
      
    pdf = FPDF(page_layout, 'mm', 'A4')
    
    #define A4 page width and height
    if page_layout == "P":
        A4_page_width = 210
        A4_page_height = 297
    else:
        A4_page_width = 297
        A4_page_height = 210

    pdf.set_auto_page_break(False)        
    img=None
    count = 0
    for id, image in tqdm(enumerate(image_list), desc="Inserting images to pdf",total=len(image_list)):
        try:
                        
            #get the width and height of the image by PIL
            
            #use PIL to open the image file based on file extension

            img = PIL.Image.open(image).convert('RGB')
            
            #get dpi information from the image
            
            img = ResizeImage(img,size=8, convert=True)    
            
            #get the width and height of the temp file image by PIL
            
            width_pixel = img.size[0]
            height_pixel = img.size[1]
            dpi = img.info.get('dpi', (300, 300))   

            #convert width and height to mm
            width_mm = int(width_pixel / dpi[0] * 25.4)            
            height_mm = int(height_pixel / dpi[1] * 25.4)            
                                            
            #save img to temp file
            temp_file = f"temp_{id}.jpg"
            img.save(temp_file, "JPEG",dpi=dpi)

            #debug print filename, width and height and dpi
            #print(f"{image}  w x h = {width_mm}x{height_mm} dpi={dpi}")
            
            pdf.add_page()            
            
            #calculate the x and y poistion of the image based on page size, make sure the image is in the center of the page  
            x_margin = 20  
            
            #calculate the y_margin based on width_mm and height_mm to make sure the image aspect ratio keeps the same
            y_margin = int (x_margin * (height_mm / width_mm))

            pos_x = int((A4_page_width- width_mm) / 2 + x_margin)
            pos_y = int((A4_page_height- height_mm) / 2 + y_margin)
            
            width_mm = width_mm -2*x_margin
            height_mm = height_mm -2*y_margin
            
            #insert the image to pdf, make sure the size of the image is not bigger than the page and keep ratio            
            pdf.image(temp_file, pos_x, pos_y, w=width_mm, h=height_mm)
            
            #remove temp file
            os.remove(temp_file)
            img.close()
            img=None
            
            count+=1
              
        except Exception as e:
            if windowGUI:
                sg.popup_error(f"[Exception]: when handling file {image}, Exception happened:\n{e}\n")
                
            else:                
                print(f"[Exception]: when handling file {image}, Exception happened:\n{e}\n")   
                
            #remove temp file if exists
            if os.path.isfile(temp_file):
                os.remove(temp_file)
            
            #close img if it's open
            if img is not None:
                img.close()
            continue
        
    pdf.output(pdf_name, "F")
    
    if windowGUI:        
        sg.popup(f"PDF file {pdf_name} created, {count} images added!")        
    else:
        print(f"PDF file {pdf_name} created, {count} images added!")   

    return count
    

def ResizeImage(image,size=6, convert=True):
    '''image is an image which is already opened by PIL.
       size is the size of the image resolution,
       convert is the option if the pic should be converted to RGB if it's RGBA
    '''
    new_height=0
    new_width=0    
    
    #first get width and height of the image
    #open the image
    
    width_pixel = image.size[0]
    height_pixel = image.size[1]
    dpi = image.info.get('dpi', (300, 300))   
      
    #calculate the new width and height
    if width_pixel > height_pixel:
        new_width = int(size * dpi[0])
        new_height = int(new_width * height_pixel / width_pixel)
    else:
        new_height = int(size * dpi[1])
        new_width = int(new_height * width_pixel / height_pixel)
    
    #print the resize info
    #print(f"\nResize image from {width_pixel}x{height_pixel} to {new_width}x{new_height}, dpi={dpi}")
    
    #resize the image
    img = image.resize((int(new_width), int(new_height)), PIL.Image.ANTIALIAS)
    
    #convert the image to RGB if it's RGBA
    if convert:
        img = img.convert('RGB')
    
    return img

if __name__ == "__main__":
    main()
    